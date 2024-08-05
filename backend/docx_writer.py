from typing import Optional
from pathlib import Path
import io

from docx import Document
from docx.text.paragraph import Paragraph

INDEX_PURPOSE_PARAGRAPH = 8
INDEX_PROCEDURE_PARAGRAPH = 10
INDEX_DURATION_PARAGRAPH = 12
INDEX_RISK_PARAGRAPH = 15
INDEX_BENEFIT_PARAGRAPHS = 17


class ICFDocument:
    def __init__(
        self,
        purpose: Optional[str] = "Not provided",
        procedure: Optional[str] = "Not provided",
        risks: Optional[str] = "Not provided",
        benefits: Optional[str] = "Not provided",
        template_fp: Optional[Path] = None,
    ):
        self._template_fp = template_fp if template_fp else "ICF-template.docx"
        self._document: Document = Document(self._template_fp)

        self.write_purpose(purpose)
        self.write_procedures(procedure)
        self.write_risks(risks)
        self.write_benefits(benefits)

    def to_stream(self, stream: io.BytesIO):
        self._document.save(stream)

    def write_purpose(self, content: str):
        self._write_paragraph_by_index(INDEX_PURPOSE_PARAGRAPH, content)

    def write_procedures(self, content: str):
        self._write_paragraph_by_index(INDEX_PROCEDURE_PARAGRAPH, content)

    def write_time_duration(self, content: str):
        self._write_paragraph_by_index(INDEX_DURATION_PARAGRAPH, content)

    def write_risks(self, content: str):
        self._write_paragraph_by_index(INDEX_RISK_PARAGRAPH, content)

    def write_benefits(self, content: str):
        self._write_paragraph_by_index(INDEX_BENEFIT_PARAGRAPHS, content)

    def _write_paragraph_by_index(self, index: int, content: str):
        self._document.paragraphs[index].text = content


if __name__ == "__main__":
    d = ICFDocument(template_fp=Path("ICF-template.docx"))
    file_stream = io.BytesIO()
    d.to_stream(file_stream)
    file_stream.seek(0)

    with open("test.docx", "wb") as f:
        f.write(file_stream.getbuffer())
